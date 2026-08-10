from docx import Document
from docx.shared import Pt

from services.storage_manager import storage

OUTPUT_DIR = storage.ensure_dir("output")

def generate_cv_docx(data: dict) -> str:
    """
    Génère un CV au format Word (.docx) basé sur les données.
    """
    doc = Document()
    
    # Style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # En-tête
    header = doc.add_heading(f"{data.get('first_name', '')} {data.get('last_name', '')}", 0)
    header.alignment = 1  # Center
    
    # Construction dynamique pour gérer l'absence d'adresse proprement
    contact_parts = [p for p in [data.get('email'), data.get('phone')] if p]
    
    location_parts = [loc_part for loc_part in [data.get('city'), data.get('residence_country')] if loc_part]
    if location_parts:
        contact_parts.append(", ".join(location_parts))
        
    contact_info = " | ".join(contact_parts)
    p = doc.add_paragraph(contact_info)
    p.alignment = 1
    if data.get('linkedin'):
        doc.add_paragraph(f"LinkedIn: {data.get('linkedin')}", style='Body Text').alignment = 1

    # Profil
    if data.get('bio'):
        doc.add_heading('Profil Professionnel', level=1)
        doc.add_paragraph(data['bio'])

    # Expérience
    if data.get('experiences'):
        doc.add_heading('Expérience Professionnelle', level=1)
        for exp in data['experiences']:
            p = doc.add_paragraph()
            runner = p.add_run(f"{exp.get('role', 'Poste')} - {exp.get('company', 'Entreprise')}")
            runner.bold = True
            p.add_run(f"\n{exp.get('start_date', '')} - {exp.get('end_date', '')}")

    # Compétences
    if data.get('skills'):
        doc.add_heading('Compétences', level=1)
        doc.add_paragraph(data['skills'])

    # Intérêts
    if data.get('interests'):
        doc.add_heading('Centres d\'intérêt', level=1)
        interests = data['interests']
        if isinstance(interests, list):
            doc.add_paragraph(", ".join(interests))
        else:
            doc.add_paragraph(str(interests))

    filename = f"cv_{data.get('last_name', 'candidate')}.docx"
    filepath = storage.path("output", filename)

    doc.save(filepath)
    return filepath